#!/usr/bin/env python3
"""
PrepMind — MCQ Extraction Script  (Free, No API)
=================================================
Extracts questions from PDFs and DOCXs using text parsing + regex.
No API key needed. Completely free.

Run:  python extract.py
"""

import re
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    import fitz  # PyMuPDF
except ImportError:
    print("❌  Run:  pip install PyMuPDF")
    sys.exit(1)

# ── Config ────────────────────────────────────────────────────────────────────

PDF_ROOT     = Path("pdfs")
OUTPUT_FILE  = Path("output/all_questions.json")
PROGRESS_DIR = Path("progress")

SUBJECT_STREAM = {
    "Physics":     "JEE",
    "Chemistry":   "BOTH",
    "Mathematics": "JEE",
    "Biology":     "NEET",
}

SKIP_FILENAMES = {"answers_physics.pdf", "answers_mathematics.pdf"}

# ── Text normalisation ────────────────────────────────────────────────────────

def normalize(text: str) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()

# ── Answer-key table parser ───────────────────────────────────────────────────

def extract_answer_key(text: str) -> dict:
    """
    Handles Biology-style separate answer tables:
      Answer key
      Question  1  2  3  4 ...
      Answer    2  2  3  3 ...
    """
    answers: dict[int, str] = {}

    m = re.search(
        r'(?:answer\s+key|ans(?:wer)?\.?\s*key|answers?)\s*\n+'
        r'(?:question\s+)?([\d\s]+)\n+'
        r'(?:answer\s+)?([\w\s]+)',
        text, re.IGNORECASE
    )
    if m:
        nums = m.group(1).split()
        vals = m.group(2).split()
        for n, v in zip(nums, vals):
            try:
                num = int(n)
                v   = v.upper()
                if v in '1234':
                    v = 'ABCD'[int(v) - 1]
                if v in 'ABCD':
                    answers[num] = v
            except (ValueError, IndexError):
                pass

    return answers

# ── Question-block splitter ───────────────────────────────────────────────────

def find_question_blocks(text: str) -> list:
    """
    Splits text into (question_number, block_text) tuples.
    Recognises: Q.1  Q. 1  1.  1)  Q1.
    """
    pat = re.compile(
        r'(?:^|\n)[ \t]*(?:Q\.?\s*)?(\d{1,3})[ \t]*[.)]\s+(?=[A-Za-z\d(])',
        re.MULTILINE
    )
    matches = list(pat.finditer(text))
    if not matches:
        return []

    blocks = []
    for i, m in enumerate(matches):
        end   = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        num   = int(m.group(1))
        block = text[m.start():end].strip()
        blocks.append((num, block))

    return blocks

# ── Option extractor ──────────────────────────────────────────────────────────

def parse_options(block: str) -> dict | None:
    """
    Extracts 4 options handling:
      (A) / (B) / (C) / (D)
      (a) / (b) / (c) / (d)
      (1) / (2) / (3) / (4)
    Returns dict with keys A B C D, or None if not found.
    """
    patterns = [
        # letter options
        (r'\(A\)', r'\(B\)', r'\(C\)', r'\(D\)'),
        (r'\(a\)', r'\(b\)', r'\(c\)', r'\(d\)'),
        # numeric options
        (r'\(1\)', r'\(2\)', r'\(3\)', r'\(4\)'),
    ]

    for A, B, C, D in patterns:
        opts = {}
        delimiters = [A, B, C, D]
        labels     = ['A', 'B', 'C', 'D']

        for i, (label, start_pat) in enumerate(zip(labels, delimiters)):
            if i + 1 < len(delimiters):
                stop = delimiters[i + 1]
            else:
                stop = r'(?:\[|\n\n|Sol\.|$)'

            m = re.search(
                rf'{start_pat}[ \t]*(.*?)(?={stop})',
                block, re.DOTALL | re.IGNORECASE
            )
            if m:
                # Take first non-empty line of the option text
                val = m.group(1).strip()
                val = re.split(r'\n', val)[0].strip()
                if val:
                    opts[label] = val

        if len(opts) == 4:
            return opts

    return None

# ── Inline answer extractor ───────────────────────────────────────────────────

def parse_answer(block: str) -> str | None:
    patterns = [
        r'\[([ABCDabcd])\]',
        r'\[([1-4])\]',
        r'(?:Ans|Answer)\.?\s*:?\s*\[?([ABCDabcd1-4])\]?',
        r'(?:correct\s+(?:answer|option))\s*:?\s*([ABCDabcd])',
    ]
    for p in patterns:
        m = re.search(p, block, re.IGNORECASE)
        if m:
            a = m.group(1).upper()
            if a in '1234':
                a = 'ABCD'[int(a) - 1]
            return a
    return None

# ── Single-block parser ───────────────────────────────────────────────────────

def parse_block(num: int, block: str) -> dict | None:
    # Strip the leading "Q.N " or "N. "
    block = re.sub(r'^[ \t]*(?:Q\.?\s*)?\d{1,3}[ \t]*[.)]\s*', '', block).strip()

    options = parse_options(block)
    if not options:
        return None

    answer = parse_answer(block)

    # Question text = everything before first option marker
    first = re.search(r'\((?:A|a|1)\)', block)
    q_text = block[:first.start()].strip() if first else block.split('\n')[0].strip()
    q_text = re.sub(r'\s+', ' ', q_text).strip()

    if len(q_text) < 8:
        return None

    return {
        "questionNumber": num,
        "question":       q_text,
        "options":        options,
        "answer":         answer,
    }

# ── Metadata helpers ──────────────────────────────────────────────────────────

def infer_subject(path: Path) -> str:
    full = str(path).lower()
    if "physics"     in full: return "Physics"
    if "chemistry"   in full: return "Chemistry"
    if "mathematics" in full: return "Mathematics"
    if "biology"     in full: return "Biology"
    return "General"


def infer_chapter(path: Path) -> str:
    name = re.sub(r'^\d+[_\-\s]', '', path.stem)
    return name.replace("-", " ").replace("_", " ").strip()


def should_skip(path: Path) -> bool:
    return path.name.lower() in SKIP_FILENAMES


def progress_file(path: Path) -> Path:
    safe = re.sub(r'[^\w]', '_', path.stem)
    return PROGRESS_DIR / f"{safe}.json"


def load_progress(path: Path) -> list | None:
    pf = progress_file(path)
    if pf.exists():
        with open(pf, encoding="utf-8") as f:
            return json.load(f)
    return None


def save_progress(path: Path, questions: list):
    PROGRESS_DIR.mkdir(exist_ok=True)
    with open(progress_file(path), "w", encoding="utf-8") as f:
        json.dump(questions, f, indent=2, ensure_ascii=False)


def tag_questions(questions: list, path: Path) -> list:
    subject = infer_subject(path)
    chapter = infer_chapter(path)
    stream  = SUBJECT_STREAM.get(subject, "BOTH")
    now     = datetime.utcnow().isoformat()
    for q in questions:
        q.setdefault("subject",    subject)
        q.setdefault("chapter",    chapter)
        q.setdefault("stream",     stream)
        q.setdefault("source",     "ALLEN")
        q.setdefault("isPYQ",      False)
        q.setdefault("difficulty", "Medium")
        q.setdefault("topic",      chapter)
        q.setdefault("createdAt",  now)
        q.pop("questionNumber", None)
    return questions

# ── PDF extractor ─────────────────────────────────────────────────────────────

def extract_pdf(path: Path) -> list:
    print(f"\n  📄  {path.name}")

    cached = load_progress(path)
    if cached is not None:
        print(f"      ✅  Already done ({len(cached)} questions) — skipping")
        return cached

    try:
        doc  = fitz.open(str(path))
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
    except Exception as e:
        print(f"      ❌  Cannot open: {e}")
        return []

    text       = normalize(text)
    answer_key = extract_answer_key(text)
    blocks     = find_question_blocks(text)

    print(f"      {len(blocks)} blocks found ...", end=" ", flush=True)

    questions = []
    for num, block in blocks:
        q = parse_block(num, block)
        if q:
            if q["answer"] is None and num in answer_key:
                q["answer"] = answer_key[num]
            questions.append(q)

    questions = tag_questions(questions, path)
    save_progress(path, questions)
    print(f"{len(questions)} questions extracted")
    return questions

# ── DOCX extractor (Biology table format) ─────────────────────────────────────

def parse_biology_answer_key(doc) -> dict:
    """Parse tab-separated answer key: 1)\td\t2)\tb\t..."""
    answers: dict[int, str] = {}
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if re.search(r'\d+\)\t[a-d]', text, re.IGNORECASE):
                    for m in re.finditer(r'(\d+)\)\t([a-dA-D])', text):
                        num = int(m.group(1))
                        ans = m.group(2).upper()
                        answers[num] = ans
    return answers


def extract_doc_biology(path: Path) -> list:
    """
    Handles Biology DOCX table format:
      Row 1: ['N.', 'question text', ...]
      Row 2: ['a)', 'Option A', 'b)', 'Option B']
      Row 3: ['c)', 'Option C', 'd)', 'Option D']
    """
    import docx
    doc = docx.Document(str(path))

    answer_key = parse_biology_answer_key(doc)
    questions  = []

    for table in doc.tables:
        rows = table.rows
        i    = 0
        while i < len(rows):
            cells = [c.text.strip() for c in rows[i].cells]

            # Detect question row: first cell looks like "1." or "12."
            num_match = re.match(r'^(\d+)\.$', cells[0]) if cells else None
            if num_match and len(cells) >= 2:
                num      = int(num_match.group(1))
                q_text   = cells[1].strip()  # question text (repeated, just take index 1)

                # Next two rows should have options
                opts: dict[str, str] = {}
                for offset in [1, 2]:
                    if i + offset < len(rows):
                        opt_cells = [c.text.strip() for c in rows[i + offset].cells]
                        # Pattern: ['a)', 'text', 'b)', 'text']
                        j = 0
                        while j < len(opt_cells) - 1:
                            label_m = re.match(r'^([a-dA-D])\)$', opt_cells[j])
                            if label_m and j + 1 < len(opt_cells):
                                label = label_m.group(1).upper()
                                opts[label] = opt_cells[j + 1]
                                j += 2
                            else:
                                j += 1

                if len(opts) == 4 and len(q_text) >= 8:
                    questions.append({
                        "questionNumber": num,
                        "question":       q_text,
                        "options":        opts,
                        "answer":         answer_key.get(num),
                    })
                i += 3  # skip question + 2 option rows
            else:
                i += 1

    return questions


# ── DOCX extractor ────────────────────────────────────────────────────────────

def extract_doc(path: Path) -> list:
    print(f"\n  📝  {path.name}")

    cached = load_progress(path)
    if cached is not None:
        print(f"      ✅  Already done ({len(cached)} questions) — skipping")
        return cached

    try:
        import docx
        doc = docx.Document(str(path))
    except ImportError:
        print("      ❌  python-docx not installed. Run: pip install python-docx")
        return []
    except Exception as e:
        print(f"      ❌  Cannot open: {e}")
        return []

    # Biology files use table format — detect by checking if tables exist with
    # cells matching "N." pattern
    is_table_format = any(
        re.match(r'^\d+\.$', row.cells[0].text.strip())
        for table in doc.tables
        for row in table.rows
        if row.cells
    )

    if is_table_format:
        print(f"      Table format detected ...", end=" ", flush=True)
        questions = extract_doc_biology(path)
    else:
        text       = normalize("\n".join(p.text for p in doc.paragraphs if p.text.strip()))
        answer_key = extract_answer_key(text)
        blocks     = find_question_blocks(text)
        print(f"      {len(blocks)} blocks found ...", end=" ", flush=True)
        questions = []
        for num, block in blocks:
            q = parse_block(num, block)
            if q:
                if q["answer"] is None and num in answer_key:
                    q["answer"] = answer_key[num]
                questions.append(q)

    questions = tag_questions(questions, path)
    save_progress(path, questions)
    print(f"{len(questions)} questions extracted")
    return questions

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print("\n🚀  PrepMind Question Extractor  (Free — No API)\n")

    OUTPUT_FILE.parent.mkdir(exist_ok=True)

    pdfs  = sorted(PDF_ROOT.rglob("*.pdf"))
    docs  = sorted(PDF_ROOT.rglob("*.doc")) + sorted(PDF_ROOT.rglob("*.docx"))
    files = pdfs + docs

    if not files:
        print(f"❌  No files found in {PDF_ROOT}/")
        sys.exit(1)

    print(f"   Found {len(pdfs)} PDF(s) and {len(docs)} DOC(s)")

    done = len([f for f in files if load_progress(f) is not None])
    if done:
        print(f"   ⏭️   {done} already done — will skip\n")

    all_questions: list = []
    stats: dict[str, int] = {}

    for f in files:
        if should_skip(f):
            print(f"\n  ⏭️   Skipping {f.name}")
            continue
        questions = extract_pdf(f) if f.suffix.lower() == ".pdf" else extract_doc(f)
        all_questions.extend(questions)
        subj = infer_subject(f)
        stats[subj] = stats.get(subj, 0) + len(questions)

    with open(OUTPUT_FILE, "w", encoding="utf-8") as out:
        json.dump(all_questions, out, indent=2, ensure_ascii=False)

    print(f"\n{'─' * 50}")
    print(f"✅  Done!")
    print(f"    Total questions : {len(all_questions)}")
    for subj, count in sorted(stats.items()):
        print(f"    {subj:<15}: {count}")
    print(f"\n    Output → {OUTPUT_FILE}")
    print(f"    Next step: python import_firestore.py\n")


if __name__ == "__main__":
    main()
